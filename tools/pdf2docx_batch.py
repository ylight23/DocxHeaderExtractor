"""Create a disposable DOCX writeback corpus from source PDFs and measure conversion quality.

The PDF remains the authority for outline extraction. This tool only creates an OOXML target and
reports whether its text can support mapping PDF-derived titles back to paragraphs.
"""

from __future__ import annotations

import argparse
import json
import re
from collections import Counter
from concurrent.futures import ProcessPoolExecutor, as_completed
from pathlib import Path

import fitz
from docx import Document
from pdf2docx import Converter


WORD = re.compile(r"\w+", re.UNICODE)


def canonical_words(text: str) -> list[str]:
    return [word.casefold() for word in WORD.findall(text)]


def title_from_key_line(line: str) -> str | None:
    if "#" not in line or line.lstrip().startswith("#"):
        return None
    return line.split("#", 1)[1].strip() or None


def key_titles(key_path: Path | None) -> list[str]:
    if key_path is None or not key_path.exists():
        return []
    return [title for line in key_path.read_text(encoding="utf-8").splitlines()
            if (title := title_from_key_line(line))]


def find_key(key_root: Path | None, stem: str) -> Path | None:
    if key_root is None:
        return None
    matches = list(key_root.rglob(f"{stem}.key"))
    return matches[0] if len(matches) == 1 else None


def measure(pdf_path: Path, docx_path: Path, key_path: Path | None) -> dict:
    with fitz.open(pdf_path) as pdf:
        pdf_text = "\n".join(page.get_text("text") for page in pdf)
        pages = pdf.page_count

    document = Document(docx_path)
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    docx_text = "\n".join(paragraphs)
    pdf_words = Counter(canonical_words(pdf_text))
    docx_words = Counter(canonical_words(docx_text))
    matched_words = sum((pdf_words & docx_words).values())
    pdf_word_count = sum(pdf_words.values())

    titles = key_titles(key_path)
    normalized_paragraphs = ["".join(canonical_words(text)) for text in paragraphs]
    missing_titles = [
        title for title in titles
        if not any("".join(canonical_words(title)) in paragraph for paragraph in normalized_paragraphs)
    ]
    resolved_titles = len(titles) - len(missing_titles)

    return {
        "pdf": str(pdf_path),
        "docx": str(docx_path),
        "key": str(key_path) if key_path else None,
        "pdfPages": pages,
        "pdfTextCharacters": len(pdf_text),
        "docxParagraphs": len(paragraphs),
        "docxTextCharacters": len(docx_text),
        "docxTables": len(document.tables),
        "canonicalWordCoverage": matched_words / pdf_word_count if pdf_word_count else None,
        "keyTitles": len(titles),
        "keyTitlesFoundInDocx": resolved_titles,
        "keyCanonicalTitleCoverage": resolved_titles / len(titles) if titles else None,
        "missingKeyTitles": missing_titles,
    }


def convert(pdf_path: Path, output_root: Path, input_root: Path, key_root: Path | None, resume: bool) -> dict:
    relative = pdf_path.relative_to(input_root).with_suffix(".docx")
    docx_path = output_root / relative
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    try:
        if resume and docx_path.exists():
            return {"status": "measured_existing", **measure(pdf_path, docx_path, find_key(key_root, pdf_path.stem))}
        converter = Converter(str(pdf_path))
        try:
            converter.convert(str(docx_path))
        finally:
            converter.close()
        return {"status": "measured", **measure(pdf_path, docx_path, find_key(key_root, pdf_path.stem))}
    except Exception as error:
        return {
            "status": "conversion_error",
            "pdf": str(pdf_path),
            "docx": str(docx_path),
            "error": f"{type(error).__name__}: {error}",
        }


def convert_job(args: tuple[Path, Path, Path, Path | None, bool]) -> dict:
    return convert(*args)


def write_manifest(output_root: Path, reports: list[dict]) -> None:
    manifest = output_root / "conversion-quality.json"
    manifest.write_text(json.dumps({"files": reports}, ensure_ascii=False, indent=2), encoding="utf-8")


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("input_root", type=Path)
    parser.add_argument("output_root", type=Path)
    parser.add_argument("--keys", type=Path)
    parser.add_argument("--only", action="append", default=[], help="PDF stem to convert; repeatable")
    parser.add_argument("--workers", type=int, default=1, help="Independent PDF conversions in parallel")
    parser.add_argument("--skip", type=int, default=0, help="Skip N sorted PDFs for resumable batches")
    parser.add_argument("--limit", type=int, help="Maximum PDFs for this batch")
    parser.add_argument("--overwrite", action="store_true", help="Reconvert existing generated DOCX files")
    args = parser.parse_args()

    pdfs = sorted(args.input_root.rglob("*.pdf"))
    if args.only:
        requested = set(args.only)
        pdfs = [pdf for pdf in pdfs if pdf.stem in requested]
    pdfs = pdfs[args.skip:]
    if args.limit is not None:
        pdfs = pdfs[:args.limit]
    if not pdfs:
        raise SystemExit("No PDFs selected.")

    prior_manifest = args.output_root / "conversion-quality.json"
    prior = []
    if prior_manifest.exists():
        try:
            prior = json.loads(prior_manifest.read_text(encoding="utf-8")).get("files", [])
        except json.JSONDecodeError:
            pass
    by_pdf = {report.get("pdf"): report for report in prior}
    jobs = [(pdf, args.output_root, args.input_root, args.keys, not args.overwrite) for pdf in pdfs]
    if args.workers > 1:
        print(f"Converting {len(jobs)} PDFs with {args.workers} workers")
        with ProcessPoolExecutor(max_workers=args.workers) as executor:
            futures = [executor.submit(convert_job, job) for job in jobs]
            for future in as_completed(futures):
                report = future.result()
                by_pdf[report["pdf"]] = report
                write_manifest(args.output_root, list(by_pdf.values()))
    else:
        for job in jobs:
            print(f"Converting {job[0].name}")
            report = convert_job(job)
            by_pdf[report["pdf"]] = report
            write_manifest(args.output_root, list(by_pdf.values()))

    print(f"Wrote {prior_manifest}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
